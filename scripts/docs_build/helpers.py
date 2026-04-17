"""Reusable styling helpers for building Bilu Store .docx files."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Brand palette (matches DESIGN.md tokens)
BRAND_PRIMARY   = RGBColor(0xFF, 0x6B, 0x35)
TEXT_PRIMARY    = RGBColor(0x0F, 0x17, 0x2A)
TEXT_SECONDARY  = RGBColor(0x47, 0x55, 0x69)
TEXT_MUTED      = RGBColor(0x94, 0xA3, 0xB8)
SURFACE_LINE    = RGBColor(0xE6, 0xE8, 0xEB)
SURFACE_RAISED  = "F7F8F9"
HEADER_FILL     = "1A1A2E"
HEADER_TEXT     = RGBColor(0xFF, 0xFF, 0xFF)
SUCCESS         = RGBColor(0x12, 0xB7, 0x6A)
DANGER          = RGBColor(0xEF, 0x44, 0x44)
WARNING         = RGBColor(0xF5, 0x9E, 0x0B)


def new_doc():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    # Body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_PRIMARY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25
    return doc


def cover(doc, title, subtitle, version, status):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("BILU STORE")
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND_PRIMARY
    r.font.bold = True

    h = doc.add_paragraph()
    rh = h.add_run(title)
    rh.font.size = Pt(34)
    rh.font.bold = True
    rh.font.color.rgb = TEXT_PRIMARY

    s = doc.add_paragraph()
    rs = s.add_run(subtitle)
    rs.font.size = Pt(14)
    rs.font.color.rgb = TEXT_SECONDARY

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(18)
    rm = meta.add_run(f"Version {version}    ·    {status}    ·    Owner: Engineering    ·    Last updated: 2026-04-17")
    rm.font.size = Pt(10)
    rm.font.color.rgb = TEXT_MUTED

    rule(doc)


def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "E6E8EB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = TEXT_PRIMARY


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = TEXT_PRIMARY


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = TEXT_SECONDARY


def para(doc, text, bold=False, color=None, size=11, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def callout(doc, label, text, color=BRAND_PRIMARY):
    """Coloured pull-quote / callout."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    rl = p.add_run(f"{label}  ")
    rl.font.bold = True
    rl.font.size = Pt(10)
    rl.font.color.rgb = color
    rt = p.add_run(text)
    rt.font.size = Pt(10)
    rt.font.color.rgb = TEXT_PRIMARY
    rt.font.italic = True


def bullets(doc, items, level=0):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        r.font.size = Pt(11)
        r.font.color.rgb = TEXT_PRIMARY


def numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        r.font.size = Pt(11)
        r.font.color.rgb = TEXT_PRIMARY


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # Background shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), SURFACE_RAISED)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = TEXT_PRIMARY


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color="E6E8EB", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.autofit = False

    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        _set_cell_bg(c, HEADER_FILL)
        _set_cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = HEADER_TEXT

    # Body
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ""
            _set_cell_bg(c, "FFFFFF" if ri % 2 else SURFACE_RAISED)
            _set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = TEXT_PRIMARY

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    # Spacing after
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def kv_table(doc, kv_pairs, label_width=2.0, value_width=4.5):
    """Two-column key/value table."""
    return table(doc, ["Field", "Value"], kv_pairs, col_widths=[label_width, value_width])


def page_break(doc):
    doc.add_page_break()
